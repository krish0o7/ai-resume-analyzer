from flask import Flask, render_template, request
import re, io

app = Flask(__name__)

# ─────────────────────────────────────────
# FILE PARSERS
# ─────────────────────────────────────────
def parse_pdf(file_bytes):
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text.strip()
    except ImportError:
        return "IMPORT_ERROR:PyPDF2"
    except Exception as e:
        return ""

def parse_docx(file_bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        # also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    except ImportError:
        return "IMPORT_ERROR:python-docx"
    except Exception as e:
        return ""

# ─────────────────────────────────────────
# MASTER LISTS
# ─────────────────────────────────────────
SKILL_LIST = [
    "python", "java", "c++", "c#", "javascript", "typescript",
    "flask", "django", "fastapi", "spring", "react", "angular", "vue",
    "node", "nodejs", "html", "css", "bootstrap", "tailwind",
    "sql", "mysql", "postgresql", "mongodb", "sqlite", "redis",
    "git", "github", "gitlab", "docker", "kubernetes", "aws",
    "azure", "gcp", "linux", "bash", "rest api", "graphql",
    "machine learning", "deep learning", "tensorflow", "pytorch",
    "data structures", "algorithms", "agile", "scrum", "ci/cd",
    "pandas", "numpy", "scikit-learn", "excel", "tableau", "power bi",
    "jira", "jenkins", "ansible", "terraform", "kafka", "spark",
    "flutter", "swift", "kotlin", "ruby", "php", "rust", "go"
]

ACTION_VERBS = [
    "developed", "built", "designed", "implemented", "created",
    "managed", "led", "optimized", "improved", "deployed",
    "automated", "integrated", "collaborated", "delivered",
    "architected", "engineered", "launched", "maintained",
    "researched", "analyzed", "reduced", "increased", "achieved",
    "spearheaded", "coordinated", "streamlined", "mentored", "executed"
]

# ─────────────────────────────────────────
# HOME
# ─────────────────────────────────────────
@app.route("/")
def home():
    return render_template("index.html",
        score=None, error=None,
        found_skills=[], missing_skills=[],
        suggestions=[], structure_feedback=[],
        breakdown={}, word_count=0,
        verbs_found=[], has_numbers=False,
        resume_text=""
    )

# ─────────────────────────────────────────
# ANALYZE
# ─────────────────────────────────────────
@app.route("/analyze", methods=["POST"])
def analyze():
    resume_text = ""

    # ── File Upload — checked FIRST ───────
    uploaded_file = request.files.get("resume_file")
    if uploaded_file and uploaded_file.filename.strip() != "":
        fname  = uploaded_file.filename.strip().lower()
        fbytes = uploaded_file.read()

        def err(msg):
            return render_template("index.html", error=msg, score=None,
                found_skills=[], missing_skills=[], suggestions=[],
                structure_feedback=[], breakdown={}, word_count=0,
                verbs_found=[], has_numbers=False, resume_text="")

        if fname.endswith(".pdf"):
            parsed = parse_pdf(fbytes)
            if parsed.startswith("IMPORT_ERROR"):
                return err("PyPDF2 is not installed. Run: pip install PyPDF2")
            if not parsed.strip():
                return err("Could not extract text from this PDF. It may be scanned/image-based. Please paste the text manually.")
            resume_text = parsed

        elif fname.endswith(".docx"):
            parsed = parse_docx(fbytes)
            if parsed.startswith("IMPORT_ERROR"):
                return err("python-docx is not installed. Run: pip install python-docx")
            if not parsed.strip():
                return err("Could not extract text from this .docx file. Please paste the text manually.")
            resume_text = parsed

        elif fname.endswith(".doc"):
            return err("Old .doc format is not supported. Please save your file as .docx or .pdf and re-upload.")

        elif fname.endswith(".txt"):
            resume_text = fbytes.decode("utf-8", errors="ignore").strip()
            if not resume_text:
                return err("The uploaded .txt file appears to be empty.")

        else:
            return err(f"Unsupported file type '{fname}'. Please upload a .pdf, .docx, or .txt file.")

    # ── Fallback to pasted text ───────────
    if not resume_text:
        resume_text = request.form.get("resume_text", "").strip()

    if not resume_text or len(resume_text.strip()) < 50:
        return render_template("index.html", error="Resume text is too short or empty. Please paste or upload a proper resume.", score=None, found_skills=[], missing_skills=[], suggestions=[], structure_feedback=[], breakdown={}, word_count=0, verbs_found=[], has_numbers=False, resume_text="")

    low        = resume_text.lower()
    words      = resume_text.split()
    word_count = len(words)

    # ════════════════════════════════════════
    # SCORING — MAX 100, HARD TO GET > 85
    # ════════════════════════════════════════

    # ── 1. SECTIONS (max 24) ─────────────
    # Must have all 6 to score well. Missing any = penalty.
    sections = {
        "Contact Info":         any(k in low for k in ["email", "phone", "mobile", "@"]),
        "Summary / Objective":  any(k in low for k in ["summary", "objective", "profile", "about me"]),
        "Education":            "education" in low,
        "Experience":           any(k in low for k in ["experience", "internship", "worked at", "work experience"]),
        "Skills":               "skills" in low,
        "Projects":             any(k in low for k in ["project", "projects"]),
    }
    section_hits  = sum(1 for v in sections.values() if v)
    section_score = section_hits * 4   # 4 pts each → max 24

    # ── 2. CONTACT QUALITY (max 12) ──────
    has_email    = bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', resume_text))
    has_phone    = bool(re.search(r'(\+?\d[\d\s\-]{8,}\d)', resume_text))
    has_linkedin = "linkedin" in low
    has_github   = "github" in low

    contact_score = (
        (4 if has_email    else 0) +
        (3 if has_phone    else 0) +
        (3 if has_linkedin else 0) +
        (2 if has_github   else 0)
    )  # max 12

    # ── 3. CONTENT QUALITY (max 24) ──────
    verbs_found    = [v for v in ACTION_VERBS if v in low]
    unique_verbs   = len(set(verbs_found))
    has_numbers    = bool(re.search(
        r'\d+\s*(%|percent|projects?|years?|months?|clients?|users?|k\b|x\b|times?)', low
    ))
    # Strict length: ideal 250–550 words
    if 250 <= word_count <= 550:
        length_score = 8
    elif 150 <= word_count < 250 or 550 < word_count <= 700:
        length_score = 4
    else:
        length_score = 0

    verb_score   = min(unique_verbs * 2, 10)   # max 10 — need 5 unique verbs to max
    number_score = 6 if has_numbers else 0      # strict: no numbers = 0

    content_score = verb_score + number_score + length_score  # max 24

    # ── 4. SKILLS DEPTH (max 28) ─────────
    found_skills   = [s for s in SKILL_LIST if s in low]
    missing_skills = [s for s in SKILL_LIST if s not in low]
    skill_count    = len(found_skills)

    # Tiered: need 10+ skills to score well
    if skill_count >= 14:
        skill_score = 28
    elif skill_count >= 10:
        skill_score = 22
    elif skill_count >= 7:
        skill_score = 16
    elif skill_count >= 4:
        skill_score = 10
    elif skill_count >= 2:
        skill_score = 5
    else:
        skill_score = 0

    # ── 5. FORMATTING PENALTIES ──────────
    penalties = 0
    # Generic filler phrases
    filler = ["responsible for", "worked on", "helped with", "assisted in", "duties included"]
    filler_hits = sum(1 for f in filler if f in low)
    penalties += filler_hits * 2  # -2 each

    # Buzzwords without evidence
    buzzwords = ["hardworking", "team player", "passionate", "go-getter", "dynamic", "synergy"]
    buzz_hits = sum(1 for b in buzzwords if b in low)
    penalties += buzz_hits * 1

    # No quantified results at all → extra penalty
    if not has_numbers:
        penalties += 5

    # Very short resume
    if word_count < 150:
        penalties += 8

    # Missing critical sections
    if not sections["Experience"]:
        penalties += 6
    if not sections["Skills"]:
        penalties += 4
    if not sections["Education"]:
        penalties += 4

    # ── FINAL SCORE ──────────────────────
    raw   = section_score + contact_score + content_score + skill_score
    raw   = max(raw - penalties, 0)
    # Hard cap: max achievable is ~88 under perfect conditions
    final_score = min(round((raw / 88) * 100, 1), 88)
    # Apply one last penalty: if < 3 verbs, further reduce
    if unique_verbs < 3:
        final_score = max(final_score - 8, 0)

    # ── STRUCTURE FEEDBACK ───────────────
    structure_feedback = [("✔" if v else "✘", k) for k, v in sections.items()]

    # ── SUGGESTIONS ──────────────────────
    suggestions = []
    if not has_email:
        suggestions.append("Add your email address — ATS systems require it for contact parsing.")
    if not has_phone:
        suggestions.append("Include your phone number for recruiter follow-up.")
    if not has_linkedin:
        suggestions.append("Add your LinkedIn profile URL — many ATS tools rank it higher.")
    if not has_github:
        suggestions.append("Add your GitHub profile link to showcase real projects.")
    if not sections["Summary / Objective"]:
        suggestions.append("Add a 2–3 line professional Summary or Objective at the very top.")
    if not sections["Projects"]:
        suggestions.append("Add a Projects section — list tech stack, your role, and impact.")
    if not has_numbers:
        suggestions.append("Quantify achievements: e.g. 'Reduced API latency by 35%', 'Built for 10k+ users'.")
    if unique_verbs < 5:
        suggestions.append(f"Only {unique_verbs} unique action verbs found. Use at least 8: built, deployed, optimized, led…")
    if word_count < 250:
        suggestions.append(f"Resume is too short ({word_count} words). Ideal is 300–550 words.")
    if word_count > 700:
        suggestions.append(f"Resume is too long ({word_count} words). ATS prefers under 700 words.")
    if skill_count < 7:
        suggestions.append(f"Only {skill_count} skills detected. Add at least 8–12 relevant technical skills.")
    if filler_hits > 0:
        suggestions.append(f"Remove passive filler phrases like 'responsible for' or 'worked on' — use action verbs instead.")
    if buzz_hits > 0:
        suggestions.append("Remove vague buzzwords like 'hardworking' or 'team player' — prove it with results instead.")
    if not sections["Education"]:
        suggestions.append("Add an Education section — ATS systems often require it to parse your profile.")
    if not suggestions:
        suggestions.append("Strong resume! Fine-tune by tailoring skills to each specific job description.")

    # ── BREAKDOWN ────────────────────────
    breakdown = {
        "Sections":        {"score": section_score,              "max": 24},
        "Contact Info":    {"score": contact_score,              "max": 12},
        "Content Quality": {"score": max(content_score - max(penalties - 5, 0), 0), "max": 24},
        "Skills Depth":    {"score": skill_score,                "max": 28},
    }

    return render_template("index.html",
        score=final_score,
        found_skills=found_skills,
        missing_skills=missing_skills[:15],
        suggestions=suggestions,
        structure_feedback=structure_feedback,
        breakdown=breakdown,
        word_count=word_count,
        verbs_found=verbs_found,
        has_numbers=has_numbers,
        resume_text=resume_text[:3000]
    )

if __name__ == "__main__":
    app.run(debug=True)